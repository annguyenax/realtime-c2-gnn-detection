"""
Chèn Mục lục, Danh sách hình, Danh sách bảng vào BaoCao_C2GNN_Final.docx.
Tìm đúng trang → xóa placeholder → chèn nội dung tĩnh đã format.
Usage:
    python scripts/fix_toc_docx.py
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent.parent
FILE = ROOT / "BaoCao_C2GNN_Final.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def _r(p, text, size=13, bold=False, italic=False, name="Times New Roman"):
    run = p.add_run(text)
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return run

def _add_tab_stop(p, position_cm):
    """Add a right-aligned dot-leader tab stop to a paragraph (1 cm = 567 twips)."""
    pPr = p._p.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs')
        pPr.append(tabs)
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), str(int(position_cm * 567)))
    tabs.append(tab)

def toc_line(doc, text, page, level=1, tab_cm=14.5):
    """One TOC entry: text ......... page"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3
    if level == 1:
        pf.left_indent = Cm(0)
        sz, bold = 13, True
    elif level == 2:
        pf.left_indent = Cm(0.5)
        sz, bold = 13, False
    else:
        pf.left_indent = Cm(1.0)
        sz, bold = 12, False
    _add_tab_stop(p, tab_cm)
    _r(p, text, size=sz, bold=bold)
    _r(p, '\t', size=sz)
    _r(p, str(page), size=sz)
    return p

def section_title(doc, text, size=13, bold=True):
    """Bold section header inside TOC pages."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    _r(p, text, size=size, bold=bold)
    return p

def fig_line(doc, num, caption, page, tab_cm=14.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3
    _add_tab_stop(p, tab_cm)
    _r(p, f"Hình {num}. {caption}", size=13)
    _r(p, '\t', size=13)
    _r(p, str(page), size=13)
    return p

def tbl_line(doc, num, caption, page, tab_cm=14.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3
    _add_tab_stop(p, tab_cm)
    _r(p, f"Bảng {num}. {caption}", size=13)
    _r(p, '\t', size=13)
    _r(p, str(page), size=13)
    return p

# ── content data ─────────────────────────────────────────────────────────────

TOC_ENTRIES = [
    # (text, page_approx, level)
    ("NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN",         "i",   1),
    ("BẢNG PHÂN CÔNG CÔNG VIỆC",                    "ii",  1),
    ("LỜI CẢM ƠN",                                  "iii", 1),
    ("TÓM TẮT ĐỀ TÀI",                              "iv",  1),
    ("MỤC LỤC",                                     "v",   1),
    ("DANH SÁCH HÌNH ẢNH",                           "vii", 1),
    ("DANH SÁCH BẢNG BIỂU",                          "viii",1),
    ("DANH MỤC TỪ VIẾT TẮT",                        "ix",  1),
    ("CHƯƠNG 1. GIỚI THIỆU",                         10,    1),
    ("1.1. Bối cảnh nghiên cứu",                     10,    2),
    ("1.2. Lý do chọn đề tài",                       11,    2),
    ("1.3. Mục tiêu đề tài",                         12,    2),
    ("1.4. Phạm vi đề tài",                          13,    2),
    ("1.5. Phương pháp nghiên cứu",                  13,    2),
    ("1.6. Bố cục báo cáo",                          14,    2),
    ("CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ NGHIÊN CỨU LIÊN QUAN", 15, 1),
    ("2.1. Mạng botnet và kênh C2",                  15,    2),
    ("2.2. Graph Neural Networks",                   17,    2),
    ("2.2.1. Tổng quan GNN và message passing",      17,    3),
    ("2.2.2. GraphSAGE (Hamilton et al., 2017)",     18,    3),
    ("2.2.3. GATv2 (Brody et al., 2022)",            18,    3),
    ("2.2.4. Tại sao GNN phù hợp bài toán C2 detection", 19, 3),
    ("2.3. Xử lý mất cân bằng lớp",                 19,    2),
    ("2.4. Dynamic graph và sliding window",         20,    2),
    ("2.5. XGBoost",                                 21,    2),
    ("2.6. Các độ đo đánh giá",                      21,    2),
    ("2.7. Nghiên cứu liên quan",                    22,    2),
    ("CHƯƠNG 3. MÔ HÌNH ĐỀ XUẤT",                   23,    1),
    ("3.1. Tổng quan hệ thống",                      23,    2),
    ("3.2. Pipeline xử lý dữ liệu (offline)",        24,    2),
    ("3.3. Xây dựng dynamic graph",                  25,    2),
    ("3.4. Thiết kế đặc trưng",                      26,    2),
    ("3.5. Các mô hình huấn luyện",                  27,    2),
    ("3.5.1. XGBoost (baseline flow-level)",         27,    3),
    ("3.5.2. GraphSAGE v3 (mô hình đề xuất)",        28,    3),
    ("3.5.3. GATv2 (baseline node-level)",           29,    3),
    ("3.6. Cơ chế threshold tuning",                 30,    2),
    ("3.7. Hệ thống phát hiện thời gian thực",       31,    2),
    ("3.7.1. Thread 1 — FlowBuilderWorker",          31,    3),
    ("3.7.2. Thread 2 — GraphUpdateWorker",          32,    3),
    ("3.7.3. Thread 3 — InferenceWorker",            32,    3),
    ("3.7.4. FastAPI Alert API và Streamlit Dashboard", 33,  3),
    ("CHƯƠNG 4. THỰC NGHIỆM, ĐÁNH GIÁ VÀ THẢO LUẬN", 34,  1),
    ("4.1. Môi trường thực nghiệm",                  34,    2),
    ("4.2. Dữ liệu thực nghiệm",                     35,    2),
    ("4.3. Kết quả huấn luyện mô hình",              36,    2),
    ("4.3.1. XGBoost",                               36,    3),
    ("4.3.2. GraphSAGE v3 — báo cáo dual-threshold", 37,    3),
    ("4.3.3. GATv2",                                 39,    3),
    ("4.3.4. Ablation: đóng góp của temporal features", 39, 3),
    ("4.4. So sánh tổng hợp các mô hình",            40,    2),
    ("4.5. Phân tích và thảo luận",                  42,    2),
    ("4.5.1. Tại sao XGBoost F1 cao hơn GNN trên CTU-13?", 42, 3),
    ("4.5.2. Cold-start gap",                        43,    3),
    ("4.5.3. Đánh giá realtime pipeline",            44,    3),
    ("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN",                45,    1),
    ("TÀI LIỆU THAM KHẢO",                           48,    1),
]

FIGURES = [
    (1,  "Bảng 18 node features đầy đủ — nhóm flow stats và temporal beaconing",             26),
    (2,  "Phân bố nhãn CTU-13 Scenario 10: 93.78% bình thường vs 6.22% botnet C2",           35),
    (3,  "SHAP top 10 features quan trọng nhất của XGBoost",                                  36),
    (4,  "F1, Precision, Recall theo threshold GraphSAGE — optimal threshold=0.9118",         38),
    (5,  "PR curve GraphSAGE (PR-AUC=0.6485)",                                                38),
    (6,  "So sánh F1, Precision, Recall, AUC-ROC bốn mô hình",                               41),
    (7,  "So sánh False Positive Rate bốn mô hình (đường đỏ: ngưỡng 0.1%)",                  41),
    (8,  "So sánh inference latency (đường đỏ: ngưỡng 100ms)",                                41),
]

TABLES = [
    (1,  "So sánh các nghiên cứu liên quan về phát hiện botnet",                              22),
    (2,  "Thống kê phân chia dữ liệu CTU-13 Scenario 10",                                    24),
    (3,  "Siêu tham số XGBoost",                                                              27),
    (4,  "Siêu tham số GraphSAGE v3",                                                         28),
    (5,  "Schema của Alert JSON object",                                                       33),
    (6,  "Môi trường thực nghiệm",                                                            34),
    (7,  "Thống kê chi tiết CTU-13 Scenario 10",                                              35),
    (8,  "Kết quả GraphSAGE v3 với hai ngưỡng quyết định",                                    37),
    (9,  "Ablation: đóng góp của 4 temporal beaconing features",                              39),
    (10, "So sánh tổng hợp bốn mô hình trên CTU-13 Scenario 10 test set",                    40),
]

# ── main patch logic ──────────────────────────────────────────────────────────

def find_section_idx(paragraphs, marker_text):
    """Return index of paragraph whose text contains marker_text (case-insensitive)."""
    needle = marker_text.upper().strip()
    for i, p in enumerate(paragraphs):
        if needle in p.text.upper().strip():
            return i
    return -1

def clear_section_content(paragraphs, start_idx):
    """
    Delete paragraphs from start_idx+1 up to (but not including) the next
    page-break paragraph or Heading-1 paragraph.
    Returns list of paragraph XML elements removed.
    """
    body = paragraphs[start_idx]._p.getparent()
    to_remove = []
    for p in paragraphs[start_idx + 1:]:
        # stop at next heading-1 or page break in a non-heading paragraph
        style = p.style.name if p.style else ""
        if style == "Heading 1":
            break
        has_page_break = any(
            br.get(qn("w:type")) == "page"
            for br in p._p.iter(qn("w:br"))
        )
        if has_page_break:
            to_remove.append(p._p)
            break
        to_remove.append(p._p)
    for el in to_remove:
        body.remove(el)

def insert_after(anchor_p, new_paragraphs_fn, doc):
    """
    Call new_paragraphs_fn(doc) which appends paragraphs to doc,
    then move those new paragraphs right after anchor_p in the XML body.
    """
    body = doc.element.body
    anchor_el = anchor_p._p

    # snapshot current last child index
    before_count = len(list(body))

    new_paragraphs_fn(doc)

    after_count = len(list(body))
    new_els = list(body)[before_count:after_count]

    # move new elements to right after anchor
    ref = anchor_el.getnext()
    for el in new_els:
        body.remove(el)
        if ref is not None:
            ref.addprevious(el)
        else:
            body.append(el)

def patch_toc(doc):
    paras = doc.paragraphs
    idx = find_section_idx(paras, "MỤC LỤC")
    if idx < 0:
        print("  WARN: MUC LUC section not found"); return
    clear_section_content(paras, idx)
    anchor = doc.paragraphs[find_section_idx(doc.paragraphs, "MỤC LỤC")]

    def write_toc(doc):
        for text, page, level in TOC_ENTRIES:
            toc_line(doc, text, page, level=level)

    insert_after(anchor, write_toc, doc)
    print("  OK  Muc luc")

def patch_figures(doc):
    paras = doc.paragraphs
    idx = find_section_idx(paras, "DANH SÁCH HÌNH ẢNH")
    if idx < 0:
        print("  WARN: DANH SACH HINH section not found"); return
    clear_section_content(paras, idx)
    anchor = doc.paragraphs[find_section_idx(doc.paragraphs, "DANH SÁCH HÌNH ẢNH")]

    def write_figs(doc):
        for num, cap, page in FIGURES:
            fig_line(doc, num, cap, page)

    insert_after(anchor, write_figs, doc)
    print("  OK  Danh sach hinh")

def patch_tables(doc):
    paras = doc.paragraphs
    idx = find_section_idx(paras, "DANH SÁCH BẢNG BIỂU")
    if idx < 0:
        print("  WARN: DANH SACH BANG section not found"); return
    clear_section_content(paras, idx)
    anchor = doc.paragraphs[find_section_idx(doc.paragraphs, "DANH SÁCH BẢNG BIỂU")]

    def write_tbls(doc):
        for num, cap, page in TABLES:
            tbl_line(doc, num, cap, page)

    insert_after(anchor, write_tbls, doc)
    print("  OK  Danh sach bang")

def main():
    if not FILE.exists():
        print(f"File not found: {FILE}")
        print("Run: python scripts/generate_report_docx.py first")
        return

    print(f"Opening {FILE.name} ...")
    doc = Document(str(FILE))

    patch_toc(doc)
    patch_figures(doc)
    patch_tables(doc)

    doc.save(str(FILE))
    size = FILE.stat().st_size // 1024
    print(f"\nSaved: {FILE.name}  ({size} KB)")
    print("Done — page numbers are approximate; finalize content then adjust if needed.")

if __name__ == "__main__":
    main()
