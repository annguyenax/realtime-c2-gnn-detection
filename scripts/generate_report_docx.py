"""
Tạo báo cáo đồ án C2GNN dạng Word (.docx).
Format dựa trên BaoCao_SSH_BruteForce_IDS.docx (lề, cỡ chữ, heading, caption, header/footer).

Usage:
    python scripts/generate_report_docx.py
Output:
    BaoCao_C2GNN_Final.docx  (trong thư mục gốc project)
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT    = Path(__file__).parent.parent
FIGURES = ROOT / "reports" / "figures"
OUT     = ROOT / "BaoCao_C2GNN_Final.docx"

# ─── Formatting helpers ──────────────────────────────────────────────────────

def _run(p, text, name="Times New Roman", size=13, bold=False, italic=False, color=None):
    run = p.add_run(text)
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def _pf(p, align=None, first_indent_cm=None, left_cm=None,
        space_before=0, space_after=6, line=1.5):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if first_indent_cm is not None:
        pf.first_line_indent = Cm(first_indent_cm)
    if left_cm is not None:
        pf.left_indent = Cm(left_cm)

def body(doc, text, first_indent=1.27):
    p = doc.add_paragraph()
    _pf(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent_cm=first_indent,
        space_before=0, space_after=3, line=1.5)
    _run(p, text, size=13)
    return p

def bullet_item(doc, text, prefix="• ", indent=1.0):
    p = doc.add_paragraph()
    _pf(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, left_cm=indent,
        space_before=0, space_after=2, line=1.5)
    _run(p, prefix + text, size=13)
    return p

def numbered_item(doc, num, text, indent=1.0):
    p = doc.add_paragraph()
    _pf(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, left_cm=indent,
        space_before=0, space_after=2, line=1.5)
    _run(p, f"{num}. {text}", size=13)
    return p

def heading1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    _pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=6, line=1.5)
    _run(p, text, size=16, bold=True)
    return p

def heading2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    _pf(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6, line=1.5)
    _run(p, text, size=14, bold=True)
    return p

def heading3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    _pf(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3, line=1.5)
    _run(p, text, size=13, bold=True, italic=True)
    return p

def caption_text(doc, text):
    p = doc.add_paragraph(style="Caption")
    _pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=8, line=1.5)
    _run(p, text, size=13, italic=True)
    return p

def blank(doc):
    p = doc.add_paragraph()
    _pf(p, space_before=0, space_after=0, line=1.0)
    return p

def centered(doc, text, size=13, bold=False):
    p = doc.add_paragraph()
    _pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=3, line=1.5)
    _run(p, text, size=size, bold=bold)
    return p

def insert_figure(doc, filename, caption_str, width_cm=14.0):
    path = FIGURES / filename
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after  = Pt(0)
    run = p_img.add_run()
    if path.exists():
        run.add_picture(str(path), width=Cm(width_cm))
    else:
        run.text = f"[HÌNH: {filename}]"
        run.font.size = Pt(11)
    caption_text(doc, caption_str)

def _field_run(p, fld_type=None, instr=None):
    """Add one run containing a single fldChar or instrText element (used in footer)."""
    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    if fld_type:
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), fld_type)
        if fld_type == "begin":
            fc.set(qn("w:dirty"), "true")
        run._r.append(fc)
    if instr:
        instrEl = OxmlElement("w:instrText")
        instrEl.set(qn("xml:space"), "preserve")
        instrEl.text = instr
        run._r.append(instrEl)
    return run

def _add_tab_stop(p, position_cm):
    """Right-aligned dot-leader tab stop (1 cm = 567 twips)."""
    pPr = p._p.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(int(position_cm * 567)))
    tabs.append(tab)

def toc_line(doc, text, page, level=1, tab_cm=14.5):
    """One TOC entry: text ......... page"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.3
    if level == 1:
        pf.left_indent = Cm(0); sz, bold = 13, True
    elif level == 2:
        pf.left_indent = Cm(0.5); sz, bold = 13, False
    else:
        pf.left_indent = Cm(1.0); sz, bold = 12, False
    _add_tab_stop(p, tab_cm)
    _run(p, text, size=sz, bold=bold)
    _run(p, "\t", size=sz)
    _run(p, str(page), size=sz)
    return p

def fig_line(doc, num, caption, page, tab_cm=14.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.3
    _add_tab_stop(p, tab_cm)
    _run(p, f"Hình {num}. {caption}", size=13)
    _run(p, "\t", size=13)
    _run(p, str(page), size=13)
    return p

def tbl_line(doc, num, caption, page, tab_cm=14.5):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.3
    _add_tab_stop(p, tab_cm)
    _run(p, f"Bảng {num}. {caption}", size=13)
    _run(p, "\t", size=13)
    _run(p, str(page), size=13)
    return p

# ── Static TOC / figure / table data ─────────────────────────────────────────

TOC_ENTRIES = [
    ("NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN",                          "i",    1),
    ("BẢNG PHÂN CÔNG CÔNG VIỆC",                                     "ii",   1),
    ("LỜI CẢM ƠN",                                                   "iii",  1),
    ("TÓM TẮT ĐỀ TÀI",                                              "iv",   1),
    ("MỤC LỤC",                                                      "v",    1),
    ("DANH SÁCH HÌNH ẢNH",                                           "vii",  1),
    ("DANH SÁCH BẢNG BIỂU",                                          "viii", 1),
    ("DANH MỤC TỪ VIẾT TẮT",                                        "ix",   1),
    ("CHƯƠNG 1. GIỚI THIỆU",                                         10,     1),
    ("1.1. Bối cảnh nghiên cứu",                                     10,     2),
    ("1.2. Lý do chọn đề tài",                                       11,     2),
    ("1.3. Mục tiêu đề tài",                                         12,     2),
    ("1.4. Phạm vi đề tài",                                          13,     2),
    ("1.5. Phương pháp nghiên cứu",                                  13,     2),
    ("1.6. Bố cục báo cáo",                                          14,     2),
    ("CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ NGHIÊN CỨU LIÊN QUAN",          15,     1),
    ("2.1. Mạng botnet và kênh C2",                                  15,     2),
    ("2.2. Graph Neural Networks",                                   17,     2),
    ("2.2.1. Tổng quan GNN và message passing",                      17,     3),
    ("2.2.2. GraphSAGE (Hamilton et al., 2017)",                     18,     3),
    ("2.2.3. GATv2 (Brody et al., 2022)",                            18,     3),
    ("2.2.4. Tại sao GNN phù hợp bài toán C2 detection",            19,     3),
    ("2.3. Xử lý mất cân bằng lớp",                                 19,     2),
    ("2.4. Dynamic graph và sliding window",                         20,     2),
    ("2.5. XGBoost",                                                 21,     2),
    ("2.6. Các độ đo đánh giá",                                      21,     2),
    ("2.7. Nghiên cứu liên quan",                                    22,     2),
    ("CHƯƠNG 3. MÔ HÌNH ĐỀ XUẤT",                                   23,     1),
    ("3.1. Tổng quan hệ thống",                                      23,     2),
    ("3.2. Pipeline xử lý dữ liệu (offline)",                        24,     2),
    ("3.3. Xây dựng dynamic graph",                                  25,     2),
    ("3.4. Thiết kế đặc trưng",                                      26,     2),
    ("3.5. Các mô hình huấn luyện",                                  27,     2),
    ("3.5.1. XGBoost (baseline flow-level)",                         27,     3),
    ("3.5.2. GraphSAGE v3 (mô hình đề xuất)",                        28,     3),
    ("3.5.3. GATv2 (baseline node-level)",                           29,     3),
    ("3.6. Cơ chế threshold tuning",                                 30,     2),
    ("3.7. Hệ thống phát hiện thời gian thực",                       31,     2),
    ("3.7.1. Thread 1 — FlowBuilderWorker",                          31,     3),
    ("3.7.2. Thread 2 — GraphUpdateWorker",                          32,     3),
    ("3.7.3. Thread 3 — InferenceWorker",                            32,     3),
    ("3.7.4. FastAPI Alert API và Streamlit Dashboard",              33,     3),
    ("CHƯƠNG 4. THỰC NGHIỆM, ĐÁNH GIÁ VÀ THẢO LUẬN",               34,     1),
    ("4.1. Môi trường thực nghiệm",                                  34,     2),
    ("4.2. Dữ liệu thực nghiệm",                                     35,     2),
    ("4.3. Kết quả huấn luyện mô hình",                              36,     2),
    ("4.3.1. XGBoost",                                               36,     3),
    ("4.3.2. GraphSAGE v3 — báo cáo dual-threshold",                 37,     3),
    ("4.3.3. GATv2",                                                 39,     3),
    ("4.3.4. Ablation: đóng góp của temporal features",              39,     3),
    ("4.4. So sánh tổng hợp các mô hình",                            40,     2),
    ("4.5. Phân tích và thảo luận",                                  42,     2),
    ("4.5.1. Tại sao XGBoost F1 cao hơn GNN trên CTU-13?",          42,     3),
    ("4.5.2. Cold-start gap",                                        43,     3),
    ("4.5.3. Đánh giá realtime pipeline",                            44,     3),
    ("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN",                                45,     1),
    ("TÀI LIỆU THAM KHẢO",                                           48,     1),
]

FIGURE_ENTRIES = [
    (1,  "Bảng 18 node features đầy đủ — nhóm flow stats và temporal beaconing",          26),
    (2,  "Phân bố nhãn CTU-13 Scenario 10: 93.78% bình thường vs 6.22% botnet C2",        35),
    (3,  "SHAP top 10 features quan trọng nhất của XGBoost",                               36),
    (4,  "F1, Precision, Recall theo threshold GraphSAGE — optimal threshold=0.9118",      38),
    (5,  "PR curve GraphSAGE (PR-AUC=0.6485)",                                             38),
    (6,  "So sánh F1, Precision, Recall, AUC-ROC bốn mô hình",                            41),
    (7,  "So sánh False Positive Rate bốn mô hình (đường đỏ: ngưỡng 0.1%)",               41),
    (8,  "So sánh inference latency (đường đỏ: ngưỡng 100ms)",                             41),
]

TABLE_ENTRIES = [
    (1,  "So sánh các nghiên cứu liên quan về phát hiện botnet",                          22),
    (2,  "Thống kê phân chia dữ liệu CTU-13 Scenario 10",                                 24),
    (3,  "Siêu tham số XGBoost",                                                           27),
    (4,  "Siêu tham số GraphSAGE v3",                                                      28),
    (5,  "Schema của Alert JSON object",                                                   33),
    (6,  "Môi trường thực nghiệm",                                                         34),
    (7,  "Thống kê chi tiết CTU-13 Scenario 10",                                           35),
    (8,  "Kết quả GraphSAGE v3 với hai ngưỡng quyết định",                                 37),
    (9,  "Ablation: đóng góp của 4 temporal beaconing features",                           39),
    (10, "So sánh tổng hợp bốn mô hình trên CTU-13 Scenario 10 test set",                 40),
]

def add_page_number_footer(doc):
    """Centered Arabic page number in footer using separate runs (required by Word)."""
    section = doc.sections[0]
    footer  = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field_run(fp, fld_type="begin")
    _field_run(fp, instr=" PAGE ")
    _field_run(fp, fld_type="separate")
    _field_run(fp, fld_type="end")

def add_table(doc, headers, rows, col_widths_cm=None, header_bg="D9E1F2"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # header row
    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = "Times New Roman"; r.font.size = Pt(12); r.font.bold = True
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), header_bg)
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = tbl.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.name = "Times New Roman"; r.font.size = Pt(12)
    if col_widths_cm:
        for ci, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return tbl

# ─── Document setup ──────────────────────────────────────────────────────────

def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin   = Cm(3.50)
    sec.right_margin  = Cm(2.01)
    sec.top_margin    = Cm(3.00)
    sec.bottom_margin = Cm(3.00)
    sec.header_distance = Cm(1.27)
    sec.footer_distance = Cm(1.27)
    add_page_number_footer(doc)
    return doc

# ─── Pages ───────────────────────────────────────────────────────────────────

def page_cover(doc):
    centered(doc, "BỘ KHOA HỌC VÀ CÔNG NGHỆ", size=14, bold=True)
    centered(doc, "HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG", size=14, bold=True)
    centered(doc, "CƠ SỞ TẠI THÀNH PHỐ HỒ CHÍ MINH", size=14, bold=True)
    centered(doc, "KHOA CÔNG NGHỆ THÔNG TIN 2", size=14, bold=True)
    centered(doc, "---o0o---", size=14, bold=True)
    blank(doc); blank(doc); blank(doc)
    centered(doc, "BÁO CÁO MÔN HỌC", size=28, bold=True)
    blank(doc)
    centered(doc,
        "Đề tài: Phát hiện luồng Command-and-Control (C2)\nbằng Graph Learning đáp ứng thời gian thực",
        size=24, bold=True)
    blank(doc); blank(doc)

    # Info table
    tbl = doc.add_table(rows=5, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER if hasattr(tbl, 'alignment') else None
    data = [
        ("Môn học", "An toàn mạng nâng cao"),
        ("GVHD", "Th.S Đàm Minh Lịnh"),
        ("Nhóm trưởng", "[Họ và tên] — [MSSV]"),
        ("Thành viên",  "[Họ và tên] — [MSSV]"),
        ("Năm học", "HK2, 2025–2026"),
    ]
    for ri, (k, v) in enumerate(data):
        for ci, val in enumerate([k, v]):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Times New Roman"
            r.font.size = Pt(13)
            r.font.bold = (ci == 0)
    blank(doc); blank(doc)
    centered(doc, "TP.HCM, tháng 06 năm 2026", size=13)
    doc.add_page_break()

def page_gvhd(doc):
    heading1(doc, "NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN")
    blank(doc)
    for _ in range(12):
        p = doc.add_paragraph()
        _pf(p, space_before=0, space_after=0, line=1.5)
        r = p.add_run("." * 90)
        r.font.name = "Times New Roman"; r.font.size = Pt(13)
    blank(doc); blank(doc)
    p = doc.add_paragraph()
    _pf(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=6, line=1.5)
    _run(p, "Giảng viên hướng dẫn\n(Ký và ghi rõ họ tên)", size=13)
    doc.add_page_break()

def page_phan_cong(doc):
    heading1(doc, "BẢNG PHÂN CÔNG CÔNG VIỆC")
    blank(doc)
    add_table(doc,
        headers=["STT", "Họ và tên — MSSV", "Nhiệm vụ"],
        rows=[
            ["1", "[Họ tên] — [MSSV]", "Trưởng nhóm — Thiết kế kiến trúc hệ thống, lập trình\nGraphSAGE/GATv2, pipeline realtime, FastAPI, viết báo cáo"],
            ["2", "[Họ tên] — [MSSV]", "Thành viên — Tiền xử lý CTU-13, huấn luyện XGBoost,\nphân tích SHAP, Streamlit dashboard, viết báo cáo"],
        ],
        col_widths_cm=[1.5, 5.5, 9.5],
    )
    doc.add_page_break()

def page_cam_on(doc):
    heading1(doc, "LỜI CẢM ƠN")
    blank(doc)
    body(doc, "Nhóm chúng em xin gửi lời cảm ơn chân thành đến Th.S Đàm Minh Lịnh — giảng viên môn An toàn mạng nâng cao tại Học viện Công nghệ Bưu chính Viễn thông cơ sở TP.HCM — đã tận tình hướng dẫn, cung cấp yêu cầu đề tài rõ ràng và hỗ trợ chúng em trong suốt quá trình thực hiện đồ án.")
    body(doc, "Chúng em cũng xin cảm ơn nhóm nghiên cứu Stratosphere IPS Lab tại Đại học Kỹ thuật Czech (CTU Prague) đã công bố tập dữ liệu CTU-13 miễn phí phục vụ nghiên cứu học thuật. Tập dữ liệu này là nền tảng thực nghiệm quan trọng giúp chúng em xây dựng và đánh giá hệ thống.")
    body(doc, "Trong quá trình thực hiện đề tài, chúng em đã cố gắng tìm hiểu, trao đổi và học hỏi từ nhiều nguồn tài liệu học thuật. Mặc dù đã nỗ lực hết mình, báo cáo vẫn không tránh khỏi những thiếu sót. Chúng em rất mong nhận được sự góp ý từ thầy và hội đồng phản biện để hoàn thiện đề tài hơn.")
    body(doc, "Cuối cùng, chúng em xin cảm ơn gia đình, bạn bè đã tạo điều kiện và động viên trong suốt quá trình học tập.")
    blank(doc)
    p = doc.add_paragraph()
    _pf(p, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=0, space_after=3, line=1.5)
    _run(p, "Xin chân thành cảm ơn.\n\nTP.HCM, ngày 13 tháng 06 năm 2026\nNhóm sinh viên thực hiện", size=13)
    doc.add_page_break()

def page_tom_tat(doc):
    heading1(doc, "TÓM TẮT ĐỀ TÀI")
    blank(doc)
    heading2(doc, "Tiếng Việt")
    body(doc, "Botnet và kênh Command-and-Control (C2) là một trong những mối đe dọa an ninh mạng nghiêm trọng nhất hiện nay. Các hệ thống phát hiện xâm nhập truyền thống dựa trên chữ ký (signature-based IDS) dễ bị vô hiệu hóa khi botnet thay đổi cổng kết nối hoặc tên miền C2. Đặc biệt, khi phân tích từng flow độc lập, các phương pháp học máy truyền thống bỏ qua cấu trúc quan hệ giữa các địa chỉ IP — đây chính là nơi ẩn chứa pattern đặc trưng của C2 beaconing.")
    body(doc, "Đề tài này trình bày một hệ thống phát hiện C2 traffic dựa trên Graph Neural Network (GNN) với Dynamic Graph Learning. Hệ thống xây dựng đồ thị động từ luồng mạng thực tế, trong đó mỗi node là một địa chỉ IP và mỗi cạnh biểu diễn một kết nối flow giữa hai IP trong cửa sổ thời gian 60 giây. Mỗi node được mô tả bằng vector đặc trưng 18 chiều bao gồm 14 thống kê flow và 4 đặc trưng temporal beaconing (active_span, mean_iat, iat_cv, repeat_dst_ratio). Pipeline thực hiện theo kiến trúc 3 luồng song song: FlowBuilderWorker, GraphUpdateWorker và InferenceWorker, đảm bảo xử lý realtime với độ trễ suy luận 18–56ms/snapshot.")
    body(doc, "Ba mô hình được huấn luyện và so sánh trên tập CTU-13 Scenario 10: XGBoost (baseline flow-level), GraphSAGE v3 (mô hình đề xuất) và GATv2 (baseline node-level). GraphSAGE với threshold tuned (0.9118) đạt F1=0.6328, Precision=0.7106, Recall=0.5703, AUC-ROC=0.9817 và FPR=0.012%. Phân tích ablation xác nhận 4 đặc trưng temporal cải thiện F1 từ 0.399 lên 0.633. Hệ thống được triển khai với FastAPI và dashboard Streamlit, tích hợp CI/CD qua GitHub Actions.")
    body(doc, "Từ khóa: botnet detection, C2 traffic, Graph Neural Network, GraphSAGE, dynamic graph, realtime IDS, CTU-13, beaconing detection.")
    blank(doc)
    heading2(doc, "English Abstract")
    body(doc, "Botnet Command-and-Control (C2) traffic poses one of the most critical cybersecurity threats. Traditional signature-based IDS systems are easily evaded when botnets rotate C2 ports or domains. Moreover, per-flow machine learning approaches ignore relational structure between IP addresses — the very structure that reveals C2 beaconing patterns.")
    body(doc, "This paper presents a realtime C2 detection system based on Dynamic Graph Neural Networks. The system builds a dynamic graph from network flows where each node represents an IP address and each edge represents a flow connection within a 60-second sliding window. Nodes are characterized by 18-dimensional feature vectors comprising 14 flow statistics and 4 temporal beaconing features (active_span, mean_iat, iat_cv, repeat_dst_ratio). A 3-thread pipeline (FlowBuilderWorker, GraphUpdateWorker, InferenceWorker) ensures realtime processing with inference latency of 18–56ms per snapshot.")
    body(doc, "Three models are trained and compared on CTU-13 Scenario 10: XGBoost (flow-level baseline), GraphSAGE v3 (proposed), and GATv2 (node-level baseline). GraphSAGE with tuned threshold (0.9118) achieves F1=0.6328, Precision=0.7106, Recall=0.5703, AUC-ROC=0.9817, and FPR=0.012%. Ablation study confirms that 4 temporal features improve F1 from 0.399 to 0.633. The system is deployed with FastAPI and Streamlit dashboard, with CI/CD via GitHub Actions.")
    body(doc, "Keywords: botnet detection, C2 traffic, Graph Neural Network, GraphSAGE, dynamic graph, realtime IDS, CTU-13, beaconing detection.")
    doc.add_page_break()

def page_muc_luc(doc):
    heading1(doc, "MỤC LỤC")
    blank(doc)
    for text, page, level in TOC_ENTRIES:
        toc_line(doc, text, page, level=level)
    doc.add_page_break()

def page_danh_sach_hinh(doc):
    heading1(doc, "DANH SÁCH HÌNH ẢNH")
    blank(doc)
    for num, cap, page in FIGURE_ENTRIES:
        fig_line(doc, num, cap, page)
    doc.add_page_break()

def page_danh_sach_bang(doc):
    heading1(doc, "DANH SÁCH BẢNG BIỂU")
    blank(doc)
    for num, cap, page in TABLE_ENTRIES:
        tbl_line(doc, num, cap, page)
    doc.add_page_break()

def page_tu_viet_tat(doc):
    heading1(doc, "DANH MỤC TỪ VIẾT TẮT")
    blank(doc)
    add_table(doc,
        headers=["Từ viết tắt", "Tiếng Anh", "Tiếng Việt"],
        rows=[
            ["C2",         "Command-and-Control",             "Kênh điều khiển botnet"],
            ["GNN",        "Graph Neural Network",            "Mạng thần kinh đồ thị"],
            ["GraphSAGE",  "Graph Sample and AggregatE",      "Mô hình GNN inductive"],
            ["GATv2",      "Graph Attention Network v2",      "Mạng chú ý đồ thị phiên bản 2"],
            ["IAT",        "Inter-Arrival Time",              "Thời gian đến giữa hai gói tin"],
            ["IDS",        "Intrusion Detection System",      "Hệ thống phát hiện xâm nhập"],
            ["FPR",        "False Positive Rate",             "Tỷ lệ cảnh báo sai"],
            ["AUC",        "Area Under Curve",                "Diện tích dưới đường cong ROC"],
            ["CTU-13",     "Czech Technical University 13",   "Tập dữ liệu botnet của CTU Prague"],
            ["IRC",        "Internet Relay Chat",             "Giao thức chat thời gian thực"],
            ["API",        "Application Programming Interface","Giao diện lập trình ứng dụng"],
            ["CI/CD",      "Continuous Integration/Delivery", "Tích hợp và triển khai liên tục"],
            ["SHAP",       "SHapley Additive exPlanations",   "Phương pháp giải thích mô hình"],
            ["PR-AUC",     "Precision-Recall AUC",            "Diện tích dưới đường PR curve"],
        ],
        col_widths_cm=[3.5, 6.0, 7.0],
    )
    doc.add_page_break()

# ─── Chapter 1 ───────────────────────────────────────────────────────────────

def chapter1(doc):
    heading1(doc, "CHƯƠNG 1. GIỚI THIỆU")
    body(doc, "Chương này trình bày bối cảnh nghiên cứu về mối đe dọa từ mạng botnet và kênh C2, lý do lựa chọn phương pháp Graph Neural Network cho bài toán phát hiện C2 traffic thời gian thực, mục tiêu, phạm vi, phương pháp nghiên cứu và bố cục báo cáo.")

    heading2(doc, "1.1. Bối cảnh nghiên cứu")
    body(doc, "Trong bối cảnh các cuộc tấn công mạng ngày càng tinh vi, mạng botnet và kênh Command-and-Control (C2) trở thành công cụ chủ chốt của tội phạm mạng. Botnet là tập hợp các máy tính đã bị nhiễm mã độc (bot), chịu sự điều khiển tập trung từ một hoặc nhiều C2 server do kẻ tấn công (botmaster) kiểm soát. Thông qua kênh C2, botmaster có thể phát động các cuộc tấn công DDoS, đánh cắp thông tin, phát tán spam hoặc triển khai ransomware trên quy mô lớn [1].")
    body(doc, "Các cuộc tấn công APT (Advanced Persistent Threat) nổi tiếng như Emotet, Mirai và Cobalt Strike đều sử dụng kênh C2 để duy trì sự hiện diện lâu dài trong hạ tầng nạn nhân. Đặc điểm nổi bật của C2 traffic là cơ chế beaconing — các bot kết nối định kỳ đến C2 server để nhận lệnh, tạo ra pattern inter-arrival time (IAT) đều đặn có thể nhận dạng được [2].")
    body(doc, "Hệ thống phát hiện xâm nhập (IDS) truyền thống dựa trên chữ ký (signature-based) đối chiếu IP/port với danh sách đen đã biết. Phương pháp này nhanh chóng bị vô hiệu hóa khi botnet sử dụng domain generation algorithm (DGA) hoặc thay đổi cổng kết nối. Mặt khác, các phương pháp học máy phân tích từng flow độc lập bỏ qua cấu trúc quan hệ giữa các địa chỉ IP — đây chính là nơi ẩn chứa đặc trưng topology của botnet (ví dụ: star pattern giữa nhiều bot và một C2 server) [3].")

    heading2(doc, "1.2. Lý do chọn đề tài")
    body(doc, "Graph-based approach cho phép mã hóa được quan hệ IP–IP: khi nhiều bot kết nối lặp lại đến cùng một C2 server, cấu trúc star topology này hiện rõ trong đồ thị mạng. Một flow đơn lẻ trông bình thường nhưng khi nhìn trong ngữ cảnh toàn đồ thị, node C2 server lộ rõ là trung tâm kết nối bất thường.")
    body(doc, "Đặc trưng beaconing (iat_cv thấp) chỉ có thể quan sát được qua nhiều flows theo thời gian. Một flow đơn không đủ thông tin, nhưng khi tích lũy trong cửa sổ thời gian 60 giây, hệ số biến thiên IAT (coefficient of variation) sẽ phân biệt rõ C2 beaconing (iat_cv ≈ 0.05–0.25) với traffic bình thường (iat_cv > 1.0) [2].")
    body(doc, "Yêu cầu thời gian thực của bài toán bảo mật đòi hỏi hệ thống phát hiện trong vòng giây đến phút, không phải giờ. Pipeline 3-thread xử lý song song đảm bảo độ trễ suy luận 56ms/snapshot — phù hợp với chu kỳ sliding window 30 giây. Tập dữ liệu CTU-13 với nhãn ground truth chính xác từ Czech Technical University là nền tảng thực nghiệm được cộng đồng học thuật tin cậy rộng rãi [1].")

    heading2(doc, "1.3. Mục tiêu đề tài")
    body(doc, "Đề tài đặt ra sáu mục tiêu cụ thể theo yêu cầu của giảng viên hướng dẫn:")
    for i, goal in enumerate([
        "Phát hiện C2 traffic dựa trên cấu trúc IP–IP và đặc trưng flow/packet, không phụ thuộc vào signature cổng hoặc IP đã biết.",
        "Tạo dynamic graph theo cửa sổ thời gian trượt (sliding window); node features là thống kê flow tổng hợp theo từng IP.",
        "Xây dựng pipeline đa luồng 3 thread: Thread 1 (FlowBuilderWorker), Thread 2 (GraphUpdateWorker), Thread 3 (InferenceWorker + alert).",
        "Huấn luyện và so sánh ba mô hình: XGBoost (baseline flow-level), GraphSAGE v3 (mô hình đề xuất) và GATv2 (baseline node-level).",
        "Đánh giá toàn diện: F1, AUC-ROC, PR-AUC, FPR, thời gian suy luận, chi phí cập nhật đồ thị.",
        "Phân tích lợi ích của graph learning so với tabular ML; đề xuất cơ chế cập nhật đồ thị nhẹ (incremental O(1)/flow) để giảm latency.",
    ], 1):
        numbered_item(doc, i, goal)

    heading2(doc, "1.4. Phạm vi đề tài")
    bullet_item(doc, "Dữ liệu: CTU-13 Scenario 10 (Murlo botnet, IRC C2, năm 2011) — proof-of-concept, không phải production dataset.")
    bullet_item(doc, "Loại C2: IRC-based beaconing qua cổng 6667. Không bao gồm HTTPS C2, domain generation algorithm (DGA), CDN fronting hay encrypted tunneling.")
    bullet_item(doc, "Nền tảng: Windows 11 Pro, CPU (không sử dụng GPU), Python 3.11.")
    bullet_item(doc, "Phạm vi đánh giá: single-scenario (Sc.10), chưa bao gồm cross-scenario evaluation.")
    bullet_item(doc, "Mục tiêu: nghiên cứu và demo — KHÔNG phải hệ thống SOC production-ready.")

    heading2(doc, "1.5. Phương pháp nghiên cứu")
    body(doc, "Đề tài áp dụng phương pháp nghiên cứu thực nghiệm gồm bốn bước: (1) Nghiên cứu tài liệu về GNN, dynamic graph learning và phát hiện botnet; (2) Phân tích dữ liệu khám phá (EDA) trên CTU-13 Scenario 10, phân tích phân bố nhãn và đặc trưng flow; (3) Thực nghiệm so sánh có kiểm soát: ba mô hình được đánh giá trên cùng tập test, cùng giao thức temporal split, cùng các độ đo; (4) Triển khai và kiểm thử hệ thống realtime với FastAPI, Streamlit và CI/CD qua GitHub Actions.")

    heading2(doc, "1.6. Bố cục báo cáo")
    body(doc, "Báo cáo được tổ chức thành bốn chương: Chương 1 trình bày bối cảnh và mục tiêu nghiên cứu. Chương 2 trình bày cơ sở lý thuyết về botnet, GNN, dynamic graph và các nghiên cứu liên quan. Chương 3 mô tả chi tiết mô hình đề xuất bao gồm kiến trúc hệ thống, pipeline xử lý dữ liệu, thiết kế đặc trưng và các mô hình học máy. Chương 4 trình bày kết quả thực nghiệm, so sánh mô hình và thảo luận chuyên sâu. Phần Kết luận tổng hợp đóng góp và đề xuất hướng phát triển.")
    doc.add_page_break()

# ─── Chapter 2 ───────────────────────────────────────────────────────────────

def chapter2(doc):
    heading1(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ NGHIÊN CỨU LIÊN QUAN")
    body(doc, "Chương này trình bày các khái niệm nền tảng cần thiết để hiểu hệ thống C2GNN: cấu trúc và hoạt động của mạng botnet, nguyên lý Graph Neural Network, phương pháp xử lý mất cân bằng lớp, cơ chế dynamic graph, mô hình XGBoost và các độ đo đánh giá. Cuối chương tóm tắt các nghiên cứu liên quan.")

    heading2(doc, "2.1. Mạng botnet và kênh C2")
    body(doc, "Mạng botnet gồm ba thành phần chính: (1) botmaster — kẻ tấn công điều khiển toàn bộ mạng; (2) C2 server — máy chủ trung gian nhận lệnh từ botmaster và phân phát đến các bot; (3) bot — các máy tính đã bị nhiễm mã độc, nhận lệnh và thực thi các tác vụ độc hại.")
    body(doc, "Giao thức C2 phổ biến bao gồm: IRC (port 6667) — đơn giản, được Murlo (CTU-13) và Neris sử dụng; HTTP/HTTPS — khó phân biệt với traffic hợp lệ; DNS tunneling — ẩn C2 trong truy vấn DNS; P2P — phi tập trung, khó triệt phá. Cơ chế beaconing là đặc trưng nhận dạng chính: các bot kết nối định kỳ đến C2 server, tạo ra inter-arrival time (IAT) đều đặn. Hệ số biến thiên IAT (iat_cv = σ/μ) của C2 beaconing thường nằm trong khoảng 0.05–0.25, trong khi traffic bình thường có iat_cv > 1.0 [1].")

    heading2(doc, "2.2. Graph Neural Networks")
    heading3(doc, "2.2.1. Tổng quan GNN và message passing")
    body(doc, "Graph Neural Network (GNN) học biểu diễn của các node trong đồ thị G = (V, E, X) thông qua framework message passing [5]. Tại mỗi bước lặp, node v cập nhật biểu diễn của mình bằng cách tổng hợp thông tin từ các neighbor:")
    body(doc, "h_v^(k) = UPDATE(h_v^(k-1), AGGREGATE({h_u^(k-1) : u ∈ N(v)}))", first_indent=2.0)
    body(doc, "trong đó N(v) là tập neighbor của v. Sau K bước lặp, biểu diễn h_v^(K) mã hóa cấu trúc cục bộ K-hop xung quanh node v.")

    heading3(doc, "2.2.2. GraphSAGE (Hamilton et al., 2017)")
    body(doc, "GraphSAGE (Graph Sample and AggregatE) [2] là phương pháp GNN inductive — có thể xử lý các node chưa thấy trong training (unseen nodes). Thay vì học embedding cho từng node cụ thể, GraphSAGE học hàm tổng hợp (aggregation function) có thể áp dụng cho bất kỳ node mới nào. Điều này đặc biệt phù hợp với bài toán C2 detection: địa chỉ IP mới xuất hiện liên tục trong traffic thực tế, không thể biết trước tất cả các node khi training.")

    heading3(doc, "2.2.3. GATv2 (Brody et al., 2022)")
    body(doc, "GATv2 [3] cải thiện Graph Attention Network (GAT) gốc [11] bằng cách sử dụng dynamic attention thay vì static attention. Trong GAT gốc, attention weight chỉ phụ thuộc vào cặp (source, target) mà không phụ thuộc vào query context — giới hạn khả năng biểu đạt. GATv2 khắc phục bằng công thức attention động: e(h_i, h_j) = a⊤ · LeakyReLU(W · [h_i ‖ h_j]), cho phép mỗi node tự học mức độ chú ý đến từng neighbor tùy theo ngữ cảnh.")

    heading3(doc, "2.2.4. Tại sao GNN phù hợp bài toán C2 detection")
    body(doc, "Phương pháp học máy truyền thống (XGBoost, SVM) phân tích từng flow độc lập — không thấy được cấu trúc quan hệ. Trong khi đó, C2 botnet tạo ra star topology đặc trưng: nhiều bot kết nối lặp lại đến cùng một node C2 server. Node C2 có fan-in cao bất thường, kết nối đến nhiều IP khác nhau với pattern IAT đều đặn. GNN mã hóa được cấu trúc này qua message passing, cho phép phân loại chính xác hơn dựa trên hành vi tổng thể của neighborhood [4].")

    heading2(doc, "2.3. Xử lý mất cân bằng lớp")
    body(doc, "CTU-13 Scenario 10 có tỷ lệ mất cân bằng nghiêm trọng: 15:1 ở mức flow và ~200:1 ở mức node sau khi xây dựng đồ thị. Mô hình huấn luyện trên dữ liệu này có xu hướng dự đoán class majority (bình thường) gần như toàn bộ thời gian, dẫn đến F1 thấp dù accuracy cao.")
    body(doc, "Các phương pháp xử lý trong hệ thống: (1) Weighted Cross-Entropy Loss với max_class_weight=50 — tăng trọng số cho class minority (botnet) trong hàm loss; (2) filter_empty=True khi training — loại bỏ các snapshot không có botnet node để model tập trung học pattern C2; (3) Threshold tuning — tìm ngưỡng quyết định tối ưu trên validation set thay vì dùng mặc định 0.5.")
    body(doc, "ROC-AUC đo khả năng ranking (threshold-independent) và phù hợp hơn accuracy khi dữ liệu mất cân bằng. PR-AUC (Precision-Recall AUC) còn nhạy cảm hơn với class minority. F1-score phụ thuộc ngưỡng nên cần báo cáo kèm threshold.")

    heading2(doc, "2.4. Dynamic graph và sliding window")
    body(doc, "Dynamic graph (đồ thị động) phản ánh sự thay đổi cấu trúc mạng theo thời gian. Hệ thống sử dụng sliding window với tham số: window_size=60s (kích thước cửa sổ), snapshot_interval=30s (tần suất chụp snapshot, 50% overlap để không bỏ sót flow biên).")
    body(doc, "Trong mỗi snapshot: node = địa chỉ IP xuất hiện trong 60s gần nhất; edge = ít nhất 1 flow giữa 2 IP trong cửa sổ; node label = botnet nếu bất kỳ adjacent edge nào là C2 flow. Incremental update O(1)/flow — thêm/xóa node và edge khi flow mới đến hoặc flow cũ hết hạn window — tránh rebuild toàn bộ đồ thị từ đầu.")

    heading2(doc, "2.5. XGBoost")
    body(doc, "XGBoost [6] (eXtreme Gradient Boosting) là thuật toán gradient boosting hiệu quả cho dữ liệu tabular. Mô hình xây dựng tuần tự các cây quyết định, mỗi cây mới khắc phục lỗi của tập cây trước đó. scale_pos_weight điều chỉnh trọng số class positive để xử lý mất cân bằng. Tính năng SHAP (SHapley Additive exPlanations) [7] cho phép giải thích đóng góp của từng feature vào quyết định cuối cùng — quan trọng cho tính minh bạch của hệ thống bảo mật.")

    heading2(doc, "2.6. Các độ đo đánh giá")
    body(doc, "Đề tài sử dụng các độ đo sau: Precision = TP/(TP+FP) — tỷ lệ cảnh báo đúng trên tổng cảnh báo; Recall = TP/(TP+FN) — tỷ lệ C2 thực sự được phát hiện; F1-score = 2·P·R/(P+R) — cân bằng Precision và Recall; AUC-ROC — diện tích dưới đường ROC, đo khả năng ranking; PR-AUC — đo hiệu năng trên class minority; FPR (False Positive Rate) = FP/(FP+TN) — tỷ lệ cảnh báo sai, rất quan trọng trong bảo mật; Latency (ms/flow hoặc ms/snapshot) — thời gian suy luận.")

    heading2(doc, "2.7. Nghiên cứu liên quan")
    body(doc, "Bảng 1 tóm tắt các nghiên cứu liên quan đến phát hiện botnet bằng GNN và học máy trên dữ liệu mạng:")
    add_table(doc,
        headers=["Tác giả, năm", "Phương pháp", "Dataset", "Kết quả"],
        rows=[
            ["Lo et al., 2022 [4]",          "E-GraphSAGE\n(GNN cho IDS)",        "CTU-13, ToN-IoT",       "F1=0.97 (CTU-13)"],
            ["Hamilton et al., 2020",        "Temporal GNN\n(fraud detection)",   "Financial graph",       "AUC=0.89"],
            ["García-Teodoro et al., 2009 [9]","Anomaly-based IDS\n(statistical)","KDD Cup 99",            "DR=99%, FAR=5.3%"],
            ["Xu et al., 2022",              "Dynamic graph\n(DynAnom)",          "Synthetic botnet",      "F1=0.81"],
            ["Đề tài này",                   "GraphSAGE v3\n(18-dim, 3-thread)",  "CTU-13 Sc.10",          "F1=0.633\nAUC=0.9817"],
        ],
        col_widths_cm=[4.0, 4.0, 3.5, 5.0],
    )
    caption_text(doc, "Bảng 1. So sánh các nghiên cứu liên quan về phát hiện botnet bằng học máy và GNN, Nguồn: tác giả tổng hợp")
    doc.add_page_break()

# ─── Chapter 3 ───────────────────────────────────────────────────────────────

def chapter3(doc):
    heading1(doc, "CHƯƠNG 3. MÔ HÌNH ĐỀ XUẤT")
    body(doc, "Chương này mô tả chi tiết kiến trúc hệ thống C2GNN: pipeline xử lý dữ liệu offline, cơ chế xây dựng dynamic graph, thiết kế đặc trưng 18 chiều, kiến trúc ba mô hình huấn luyện, cơ chế threshold tuning và hệ thống phát hiện thời gian thực.")

    heading2(doc, "3.1. Tổng quan hệ thống")
    body(doc, "Hệ thống C2GNN được thiết kế theo kiến trúc hai giai đoạn: (1) Offline training — tiền xử lý CTU-13, xây dựng đồ thị tĩnh, huấn luyện mô hình; (2) Online inference — pipeline 3-thread xử lý flow thực tế, cập nhật đồ thị động và suy luận realtime. Luồng dữ liệu tổng quát: raw flows → dynamic graph → GNN inference → alert → FastAPI → Streamlit dashboard.")
    body(doc, "Stack công nghệ: Python 3.11, PyTorch 2.3, PyTorch Geometric 2.5, XGBoost 2.x, FastAPI + Uvicorn, Streamlit, NetworkX, MLflow, Docker, GitHub Actions (CI/CD: ruff + pytest + Bandit + Trivy).")

    heading2(doc, "3.2. Pipeline xử lý dữ liệu (offline)")
    body(doc, "Dữ liệu đầu vào là file .binetflow từ CTU-13 Scenario 10 — định dạng bidirectional Argus flows, KHÔNG phải NetFlow v5/v9 tiêu chuẩn. Các bước tiền xử lý:")
    numbered_item(doc, 1, "Parse .binetflow → pandas DataFrame, giữ các cột: StartTime, Duration, Proto, SrcAddr, Sport, DstAddr, Dport, TotBytes, TotPkts, Label.")
    numbered_item(doc, 2, 'Label mapping: "Botnet" → 1; "Background", "Normal", "LEGITIMATE" → 0.')
    numbered_item(doc, 3, "Temporal split 70/15/15 (KHÔNG shuffle): train trên phần đầu, validate trên phần giữa, test trên phần cuối của chuỗi thời gian — tránh data leakage theo thời gian.")
    numbered_item(doc, 4, "Xây dựng graph snapshots: tổng 421 snapshots (train: 295, val: 63, test: 63). filter_empty=True khi training — loại snapshot không có botnet node.")
    blank(doc)
    add_table(doc,
        headers=["Tập", "Số flows", "Số snapshots", "Tỷ lệ botnet flow"],
        rows=[
            ["Train (70%)",      "≈3.624.892", "295", "6.22%"],
            ["Validation (15%)", "≈776.763",   "63",  "6.22%"],
            ["Test (15%)",       "≈776.762",   "63",  "6.22%"],
            ["Tổng",             "5.178.417",  "421", "6.22%"],
        ],
        col_widths_cm=[3.5, 4.0, 4.0, 5.0],
    )
    caption_text(doc, "Bảng 2. Thống kê phân chia dữ liệu CTU-13 Scenario 10, Nguồn: tác giả")

    heading2(doc, "3.3. Xây dựng dynamic graph")
    body(doc, "Lớp SlidingWindowGraph thực hiện cập nhật đồ thị gia tăng: với mỗi flow mới (src_ip, dst_ip, timestamp), thêm node nếu chưa tồn tại và thêm hoặc cập nhật edge. Khi timestamp - node.first_seen > window_size, loại node đó khỏi đồ thị. Chi phí cập nhật là O(1)/flow — không rebuild toàn bộ đồ thị.")
    body(doc, "Cứ mỗi snapshot_interval=30 giây, chụp trạng thái hiện tại của đồ thị dưới dạng PyG Data object (torch_geometric.data.Data). Node label được gán: botnet=1 nếu bất kỳ adjacent edge nào có nhãn C2 flow. Điều này có thể dẫn đến over-labeling cho hub nodes (DNS server, default gateway) — là hạn chế đã biết của phương pháp.")

    heading2(doc, "3.4. Thiết kế đặc trưng")
    body(doc, "Vector đặc trưng node 18 chiều được chia thành hai nhóm: 14 flow statistics và 4 temporal beaconing features (mới trong phiên bản v3).")
    blank(doc)
    insert_figure(doc, "node_features_table.png",
        "Hình 1. Bảng 18 node features đầy đủ — nhóm flow stats (xanh lam) và temporal beaconing (xanh nhạt), Nguồn: tác giả")
    body(doc, "Feature quan trọng nhất là iat_cv (coefficient of variation IAT): C2 beaconing tạo ra các kết nối định kỳ đều đặn dẫn đến σ/μ nhỏ (≈0.05–0.25), trong khi traffic người dùng bình thường ngẫu nhiên hơn (iat_cv > 1.0). Ablation study xác nhận 4 temporal features cải thiện F1 từ 0.399 (14-dim) lên 0.633 (18-dim) — tương đương +58.6% tương đối.")

    heading2(doc, "3.5. Các mô hình huấn luyện")
    heading3(doc, "3.5.1. XGBoost (baseline flow-level)")
    body(doc, "XGBoost được huấn luyện trên 14 tabular features ở mức flow — mỗi flow là một sample. Đây là baseline mạnh để so sánh: nhanh (2.1ms/flow), dễ giải thích qua SHAP, không cần xây dựng đồ thị.")
    add_table(doc,
        headers=["Siêu tham số", "Giá trị", "Lý do"],
        rows=[
            ["n_estimators",      "400",   "Đủ cây để hội tụ mà không overfit"],
            ["max_depth",         "7",     "Cân bằng độ phức tạp và khái quát hóa"],
            ["learning_rate",     "0.1",   "Mặc định hiệu quả"],
            ["scale_pos_weight",  "≈15",   "= n_negative/n_positive, xử lý 15:1 imbalance"],
            ["subsample",         "0.8",   "Giảm overfitting qua row sampling"],
            ["colsample_bytree",  "0.8",   "Giảm overfitting qua column sampling"],
        ],
        col_widths_cm=[4.0, 3.0, 9.5],
    )
    caption_text(doc, "Bảng 3. Siêu tham số XGBoost, Nguồn: tác giả")

    heading3(doc, "3.5.2. GraphSAGE v3 (mô hình đề xuất)")
    body(doc, "GraphSAGE v3 là mô hình chính của đề tài, thực hiện phân loại node (IP-level) dựa trên đồ thị động. Kiến trúc 3 lớp SAGEConv:")
    body(doc, "SAGEConv(18→128) + BatchNorm + ReLU + Dropout(0.3)\n→ SAGEConv(128→128) + BatchNorm + ReLU + Dropout(0.3)\n→ SAGEConv(128→64) + BatchNorm + ReLU\n→ Linear(64→1) + Sigmoid", first_indent=2.0)
    add_table(doc,
        headers=["Siêu tham số", "Giá trị", "Lý do"],
        rows=[
            ["Input dim",          "18",        "18-dim node features"],
            ["Hidden dim",         "128",        "Đủ sức biểu đạt, không quá nặng"],
            ["Loss",               "WeightedCE (cap=50)", "Xử lý ~200:1 node-level imbalance"],
            ["Optimizer",          "Adam, lr=0.001", "Hiệu quả với sparse gradients"],
            ["Scheduler",          "CosineAnnealingLR T_max=100", "Giữ LR ổn định, tránh diverge"],
            ["Epochs",             "50, patience=12", "Early stopping tránh overfit"],
            ["Seed",               "42",         "Reproducibility"],
            ["Learning paradigm",  "Inductive",  "Xử lý IP mới không thấy khi training"],
        ],
        col_widths_cm=[4.0, 4.5, 8.0],
    )
    caption_text(doc, "Bảng 4. Siêu tham số GraphSAGE v3, Nguồn: tác giả")

    heading3(doc, "3.5.3. GATv2 (baseline node-level)")
    body(doc, "GATv2 được dùng như baseline node-level để so sánh với GraphSAGE. Kiến trúc: 2 × GATv2Conv(heads=4, hidden=64) → Linear(64→1) + Sigmoid. Chưa thực hiện threshold tuning cho GATv2 — dùng default threshold=0.5 để thấy tác động của imbalance.")

    heading2(doc, "3.6. Cơ chế threshold tuning")
    body(doc, "Do mất cân bằng node-level ~200:1, model calibration bị shift về phía output thấp — F1 với threshold=0.5 rất thấp (0.3951) dù AUC cao (0.9817). Threshold tuning là bước bắt buộc:")
    numbered_item(doc, 1, "Sweep threshold t ∈ [0.01, 0.99] với bước 0.001 trên validation set.")
    numbered_item(doc, 2, "Tính F1(t) với constraint min_recall ≥ 0.40 (đảm bảo không bỏ sót quá nhiều C2).")
    numbered_item(doc, 3, "Chọn t* = argmax F1(t) → t* = 0.9118.")
    numbered_item(doc, 4, "Đánh giá trên test set với t* = 0.9118 — không có data leakage.")
    body(doc, "Threshold cao (0.9118) là hiện tượng tự nhiên khi imbalance nặng: model cần rất tự tin (>91%) mới báo botnet, để giảm false positive từ majority class. FPR giảm 9× từ 0.107% (default) xuống 0.012% (tuned).")

    heading2(doc, "3.7. Hệ thống phát hiện thời gian thực")
    heading3(doc, "3.7.1. Thread 1 — FlowBuilderWorker")
    body(doc, "Đọc file .parquet theo thứ tự thời gian, parse từng FlowRecord (src_ip, dst_ip, timestamp, label, features) và đẩy vào flow_queue. Tham số --realtime-factor điều chỉnh tốc độ replay: realtime-factor=50 tương đương replay nhanh hơn thực tế 50×.")

    heading3(doc, "3.7.2. Thread 2 — GraphUpdateWorker")
    body(doc, "Nhận FlowRecord từ flow_queue, cập nhật incremental graph O(1)/flow. Sau mỗi snapshot_interval=30s, tính 18 node features cho tất cả node hiện tại và đẩy PyG Data object vào graph_queue.")

    heading3(doc, "3.7.3. Thread 3 — InferenceWorker")
    body(doc, "Nhận graph snapshot từ graph_queue, thực hiện GNN forward pass, lấy P(botnet) cho mỗi node. Với node có P > threshold, tạo Alert object với thông tin: timestamp_iso, src_ip, dst_ip, risk_score, reasons (beaconing patterns), inference_latency_ms, graph_nodes, graph_edges. Alert được POST đến FastAPI endpoint /alerts và ghi vào alerts.jsonl.")

    heading3(doc, "3.7.4. FastAPI Alert API và Streamlit Dashboard")
    body(doc, "FastAPI cung cấp 4 endpoint: POST /alerts (nhận alert từ InferenceWorker), GET /alerts (lấy danh sách alerts), GET /stats (thống kê tổng hợp), GET /health (kiểm tra trạng thái). Streamlit dashboard hiển thị realtime: tổng alerts, FPR%, inference latency, bảng alerts có thể lọc, biểu đồ timeline và graph visualization với màu sắc theo botnet probability.")
    add_table(doc,
        headers=["Trường", "Kiểu", "Mô tả"],
        rows=[
            ["timestamp_iso",       "string",  "ISO 8601 timestamp phát hiện"],
            ["src_ip",              "string",  "Địa chỉ IP nguồn đáng ngờ"],
            ["dst_ip",              "string",  "Địa chỉ IP đích"],
            ["risk_score",          "float",   "P(botnet) từ GNN, range [0,1]"],
            ["reasons",             "list",    "Các pattern phát hiện (beaconing, ...)"],
            ["inference_latency_ms","float",   "Thời gian suy luận tính bằng ms"],
            ["graph_nodes",         "int",     "Số node trong snapshot"],
            ["graph_edges",         "int",     "Số cạnh trong snapshot"],
        ],
        col_widths_cm=[4.5, 2.5, 9.5],
    )
    caption_text(doc, "Bảng 5. Schema của Alert JSON object, Nguồn: tác giả")
    doc.add_page_break()

# ─── Chapter 4 ───────────────────────────────────────────────────────────────

def chapter4(doc):
    heading1(doc, "CHƯƠNG 4. THỰC NGHIỆM, ĐÁNH GIÁ VÀ THẢO LUẬN")
    body(doc, "Chương này trình bày thiết lập môi trường thực nghiệm, đặc điểm dữ liệu CTU-13 Scenario 10, kết quả huấn luyện ba mô hình với đầy đủ số liệu, so sánh tổng hợp và thảo luận chuyên sâu về các kết quả thu được.")

    heading2(doc, "4.1. Môi trường thực nghiệm")
    add_table(doc,
        headers=["Thành phần", "Cấu hình"],
        rows=[
            ["Hệ điều hành",     "Windows 11 Pro"],
            ["CPU",              "[Model CPU] (không sử dụng GPU)"],
            ["RAM",              "[Dung lượng GB]"],
            ["Python",           "3.11"],
            ["PyTorch",          "2.3"],
            ["PyTorch Geometric","2.5"],
            ["XGBoost",          "2.x"],
            ["FastAPI",          "0.111.x + Uvicorn"],
            ["Streamlit",        "1.35.x"],
        ],
        col_widths_cm=[5.0, 11.5],
    )
    caption_text(doc, "Bảng 6. Môi trường thực nghiệm, Nguồn: tác giả")

    heading2(doc, "4.2. Dữ liệu thực nghiệm")
    body(doc, "CTU-13 Scenario 10 chứa traffic của botnet Murlo — loại botnet IRC C2 sử dụng cổng 6667 với beaconing interval ~30–45 giây. Dataset được thu thập tại Czech Technical University năm 2011, bao gồm traffic từ mạng thực của trường đại học [1].")
    insert_figure(doc, "dataset_distribution.png",
        "Hình 2. Phân bố nhãn CTU-13 Scenario 10: 93.78% bình thường vs 6.22% botnet C2 (flow-level), Nguồn: tác giả")
    add_table(doc,
        headers=["Thống kê", "Giá trị"],
        rows=[
            ["Tổng flows",              "5.178.417"],
            ["Botnet flows",            "322.158 (6,22%)"],
            ["Bình thường",             "4.856.259 (93,78%)"],
            ["Imbalance (flow)",        "≈15:1"],
            ["Imbalance (node-level)",  "≈200:1"],
            ["Giao thức botnet",        "IRC, port 6667"],
            ["Beaconing interval",      "~30–45 giây"],
            ["Định dạng",               "Bidirectional Argus flows (.binetflow)"],
        ],
        col_widths_cm=[6.0, 10.5],
    )
    caption_text(doc, "Bảng 7. Thống kê chi tiết CTU-13 Scenario 10, Nguồn: Garcia et al. (2014)")

    heading2(doc, "4.3. Kết quả huấn luyện mô hình")
    heading3(doc, "4.3.1. XGBoost")
    body(doc, "XGBoost đạt kết quả xuất sắc: F1=0.9921, Precision=0.9895, Recall=0.9947, AUC-ROC=0.9998, FPR=0.10%, latency=2.1ms/flow. Đây là kết quả cao nhất trong ba mô hình về F1 và AUC.")
    body(doc, "Phân tích SHAP cho thấy src_port (SHAP mean abs=2.811) và dst_port (1.997) là hai feature có tầm ảnh hưởng lớn nhất — cách biệt xa so với các feature tiếp theo (bytes_per_packet=1.997, is_tcp=1.416, total_bytes=1.254).")
    insert_figure(doc, "shap_importance.png",
        "Hình 3. SHAP top 10 features quan trọng nhất của XGBoost — src_port và dst_port dẫn đầu, Nguồn: tác giả")

    heading3(doc, "4.3.2. GraphSAGE v3 — báo cáo dual-threshold")
    body(doc, "GraphSAGE được báo cáo với hai threshold để minh bạch về tác động của class imbalance:")
    add_table(doc,
        headers=["Threshold", "F1", "Precision", "Recall", "AUC-ROC", "FPR"],
        rows=[
            ["Default (0.5)",  "0.3951", "0.2675", "0.7557", "0.9817", "0.107%"],
            ["Tuned (0.9118)", "0.6328", "0.7106", "0.5703", "0.9817", "0.012%"],
        ],
        col_widths_cm=[3.5, 2.5, 3.0, 2.5, 3.0, 2.0],
    )
    caption_text(doc, "Bảng 8. Kết quả GraphSAGE v3 với hai ngưỡng quyết định, Nguồn: tác giả")
    body(doc, "AUC-ROC=0.9817 cho thấy model có khả năng ranking rất tốt (threshold-independent). F1 thấp với threshold=0.5 là hệ quả của class imbalance node-level ~200:1 dẫn đến model calibration bị shift — không phải model kém. Threshold tuning trên validation set đưa F1 lên 0.6328 và đặc biệt giảm FPR 9× từ 0.107% xuống 0.012%.")
    insert_figure(doc, "threshold_sweep_graphsage.png",
        "Hình 4. F1, Precision, Recall theo threshold GraphSAGE — optimal threshold=0.9118, Nguồn: tác giả")
    insert_figure(doc, "pr_curve_graphsage.png",
        "Hình 5. PR curve GraphSAGE (PR-AUC=0.6485), Nguồn: tác giả")

    heading3(doc, "4.3.3. GATv2")
    body(doc, "GATv2 với default threshold=0.5: F1=0.0518, Precision=0.0267, Recall=0.8389, AUC-ROC=0.9701, FPR=1.537%, latency=296.5ms/graph. AUC=0.9701 cho thấy model có khả năng phân biệt tốt, nhưng FPR=1.537% quá cao cho ứng dụng thực tế (cứ 1000 node bình thường thì 15 node bị cảnh báo sai). GATv2 cần threshold tuning tương tự GraphSAGE để sử dụng thực tế.")

    heading3(doc, "4.3.4. Ablation: đóng góp của temporal features")
    body(doc, "So sánh hai phiên bản GraphSAGE trên cùng tập test:")
    add_table(doc,
        headers=["Phiên bản", "Input dim", "F1 (tuned)", "Đặc trưng thêm"],
        rows=[
            ["GraphSAGE v2", "14-dim", "≈0.399", "Chỉ flow stats"],
            ["GraphSAGE v3", "18-dim", "0.6328", "+ active_span, mean_iat, iat_cv, repeat_dst_ratio"],
        ],
        col_widths_cm=[4.0, 3.0, 3.0, 6.5],
    )
    caption_text(doc, "Bảng 9. Ablation: đóng góp của 4 temporal beaconing features, Nguồn: tác giả")
    body(doc, "Improvement: +58.6% F1 tương đối (0.399→0.633). Feature iat_cv là feature quyết định: nó định lượng hóa được pattern beaconing đều đặn của C2 botnet — điều không thể làm được chỉ với flow stats tĩnh.")

    heading2(doc, "4.4. So sánh tổng hợp các mô hình")
    add_table(doc,
        headers=["Mô hình", "Threshold", "F1", "Precision", "Recall", "AUC-ROC", "FPR", "Latency"],
        rows=[
            ["XGBoost",            "0.5",    "0.9921", "0.9895", "0.9947", "0.9998", "0.10%",  "2.1ms/flow"],
            ["GraphSAGE (default)","0.5",    "0.3951", "0.2675", "0.7557", "0.9817", "0.107%", "56ms/graph"],
            ["GraphSAGE (tuned)",  "0.9118", "0.6328", "0.7106", "0.5703", "0.9817", "0.012%", "56ms/graph"],
            ["GATv2 (default)",    "0.5",    "0.0518", "0.0267", "0.8389", "0.9701", "1.537%", "297ms/graph"],
        ],
        col_widths_cm=[3.8, 2.5, 2.0, 2.5, 2.0, 2.5, 2.0, 3.2],
    )
    caption_text(doc, "Bảng 10. So sánh tổng hợp bốn mô hình trên CTU-13 Scenario 10 test set, Nguồn: tác giả")
    insert_figure(doc, "model_comparison_bar.png",
        "Hình 6. So sánh F1, Precision, Recall, AUC-ROC bốn mô hình, Nguồn: tác giả")
    insert_figure(doc, "fpr_comparison.png",
        "Hình 7. So sánh False Positive Rate bốn mô hình (đường đỏ: ngưỡng 0.1%), Nguồn: tác giả")
    insert_figure(doc, "latency_comparison.png",
        "Hình 8. So sánh inference latency (đường đỏ: ngưỡng 100ms), Nguồn: tác giả")

    heading2(doc, "4.5. Phân tích và thảo luận")
    heading3(doc, "4.5.1. Tại sao XGBoost F1 cao hơn GNN trên CTU-13?")
    body(doc, "XGBoost đạt F1=0.9921 trong khi GraphSAGE tuned chỉ đạt 0.6328. Kết quả này có vẻ mâu thuẫn với kỳ vọng về GNN, nhưng có giải thích rõ ràng: CTU-13 Scenario 10 sử dụng Murlo botnet với IRC C2 cố định qua port 6667 — một port rất đặc trưng, hiếm thấy trong traffic bình thường. SHAP phân tích cho thấy src_port (2.811) và dst_port (1.997) là hai feature tầm quan trọng cao nhất, cách biệt xa các feature khác. Điều này cho thấy XGBoost thực chất đang học một dạng signature-based detection ẩn trong framework tabular ML.")
    body(doc, "Hệ quả quan trọng: nếu botnet sử dụng port ngẫu nhiên hoặc HTTPS (port 443), XGBoost sẽ bị bypass hoàn toàn vì mất feature discriminative chính. GNN, ngược lại, học topology pattern (star structure, beaconing via iat_cv) — ít phụ thuộc port hơn và có khả năng generalize tốt hơn sang botnet dùng giao thức khác.")

    heading3(doc, "4.5.2. Cold-start gap")
    body(doc, "Warm-start F1=0.652 (graph đã có history) vs Cold-start F1=0.06 (fresh graph, không có history). Nguyên nhân: features iat_cv và repeat_dst_ratio cần ít nhất 2 flows để tính toán — với fresh graph, tất cả temporal features bằng 0, model không có đủ thông tin để phân biệt C2. Hàm ý thực tế: cần warm-up period 60–120 giây trước khi bắt đầu tạo alert chính thức. Trong production, có thể giải quyết bằng warm graph initialization từ history log của session trước.")

    heading3(doc, "4.5.3. Đánh giá realtime pipeline")
    body(doc, "Demo thực tế với 3.000 flows: 7 snapshots được tạo, 7 lượt inference, 53 alerts được POST đến FastAPI. Throughput: 3.000 flows trong 2.2 giây ≈ 1.363 flows/giây. Inference latency mẫu: 18.38ms/snapshot (tốt hơn nhiều so với 56ms trung bình do snapshot nhỏ hơn). End-to-end latency: 30–60 giây (phần lớn là thời gian tích lũy window 30 giây).")
    body(doc, "Ví dụ alert thực tế từ pipeline:")
    p = doc.add_paragraph()
    _pf(p, align=WD_ALIGN_PARAGRAPH.LEFT, left_cm=1.5, space_before=3, space_after=3, line=1.2)
    _run(p, '{ "timestamp_iso": "2011-08-18T07:13:14Z", "src_ip": "147.32.84.25",\n  "dst_ip": "195.113.232.96", "risk_score": 0.3063,\n  "reasons": ["short-lived connections (possible beaconing)"],\n  "inference_latency_ms": 18.38, "graph_nodes": 437, "graph_edges": 710 }',
         name="Courier New", size=10)
    doc.add_page_break()

# ─── Kết luận ────────────────────────────────────────────────────────────────

def ket_luan(doc):
    heading1(doc, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")

    heading2(doc, "Kết quả đạt được")
    body(doc, "Đề tài đã hoàn thành đầy đủ sáu mục tiêu đề ra:")
    for text in [
        "Xây dựng thành công pipeline 3-thread realtime end-to-end: FlowBuilderWorker → GraphUpdateWorker → InferenceWorker với incremental graph update O(1)/flow.",
        "Thiết kế dynamic graph với 18-dim node features, trong đó 4 temporal beaconing features (iat_cv, repeat_dst_ratio, active_span, mean_iat) cải thiện F1 từ 0.399 lên 0.633 (+58.6% tương đối).",
        "GraphSAGE v3 với tuned threshold (0.9118) đạt: F1=0.6328, Precision=0.7106, Recall=0.5703, AUC-ROC=0.9817, FPR=0.012%, latency=56ms/graph — đáp ứng tất cả thời hạn yêu cầu.",
        "So sánh đầy đủ và công bằng ba mô hình: XGBoost (F1=0.9921), GraphSAGE (F1=0.6328 tuned), GATv2 (F1=0.0518 default) với cùng tập test và giao thức đánh giá.",
        "Phân tích minh bạch dual-threshold cho GraphSAGE; thừa nhận cold-start gap và port-dependent limitation của XGBoost.",
        "Triển khai FastAPI + Streamlit dashboard hoạt động, sinh 53 alerts thực tế, tích hợp CI/CD qua GitHub Actions.",
    ]:
        bullet_item(doc, text)

    heading2(doc, "Hạn chế")
    for text in [
        "CTU-13 năm 2011 không đại diện C2 hiện đại: HTTPS C2, DGA domain rotation, CDN fronting và P2P C2 không có trong dữ liệu.",
        "Single-scenario: chưa validate cross-scenario (train Sc.10 → test Sc.8 Rbot) — khả năng generalize chưa được kiểm chứng.",
        "Cold-start gap nghiêm trọng: F1=0.652 (warm) vs F1=0.06 (cold) — cần warm-up 60–120 giây trước khi hoạt động.",
        "XGBoost F1 cao do port-dependent detection (IRC port 6667), không phải vì khả năng phát hiện behavior chung.",
        "NetworkX backend không scale cho >10k concurrent nodes; cần graph database (GraphBolt, DGL distributed) cho production.",
    ]:
        bullet_item(doc, text)

    heading2(doc, "Hướng phát triển")
    for text in [
        "Cross-scenario evaluation: train trên Scenario 10 (Murlo IRC), test trên Scenario 8 (Rbot HTTP) để đánh giá khả năng generalize.",
        "Temporal GNN (TGN — Temporal Graph Network) cho time-aware graph evolution, tích hợp trực tiếp lịch sử thời gian vào kiến trúc model.",
        "Warm graph initialization: lưu graph state giữa các session để giảm cold-start gap xuống gần 0.",
        "Mở rộng sang C2 hiện đại: thu thập dữ liệu HTTPS C2, DGA, thêm TLS metadata features (JA3 fingerprint, SNI entropy).",
        "Production-scale streaming: thay NetworkX bằng Kafka + Flink, sử dụng distributed graph processing (GraphBolt, DGL distributed).",
        "Dataset bổ sung: CICIDS2017 [12], UNSW-NB15 để tăng diversity và validate cross-dataset generalization.",
    ]:
        bullet_item(doc, text)
    doc.add_page_break()

# ─── References ──────────────────────────────────────────────────────────────

def tai_lieu_tham_khao(doc):
    heading1(doc, "TÀI LIỆU THAM KHẢO")
    refs = [
        '[1] S. Garcia, M. Grill, J. Stiborek, and A. Zunino, "An empirical comparison of botnet detection methods," Computers & Security, vol. 45, pp. 100-123, 2014.',
        '[2] W. L. Hamilton, Z. Ying, and J. Leskovec, "Inductive Representation Learning on Large Graphs," in Proc. NeurIPS, 2017, pp. 1024-1034.',
        '[3] S. Brody, U. Alon, and E. Yahav, "How Attentive are Graph Attention Networks?" in Proc. ICLR, 2022.',
        '[4] A. Lo, H. Layeghy, M. Sarhan, M. Gallagher, and M. Portmann, "E-GraphSAGE: A Graph Neural Network Based Intrusion Detection System for IoT," in Proc. IEEE/IFIP NOMS, 2022.',
        '[5] T. N. Kipf and M. Welling, "Semi-Supervised Classification with Graph Convolutional Networks," in Proc. ICLR, 2017.',
        '[6] T. Chen and C. Guestrin, "XGBoost: A Scalable Tree Boosting System," in Proc. ACM SIGKDD, 2016, pp. 785-794.',
        '[7] S. M. Lundberg and S.-I. Lee, "A Unified Approach to Interpreting Model Predictions," in Proc. NeurIPS, 2017.',
        '[8] Stratosphere IPS Lab, "CTU-13 Botnet Dataset," [Online]. Available: https://www.stratosphereips.org/datasets-ctu13. [Accessed: Jun. 2026].',
        '[9] P. Garcia-Teodoro, J. Diaz-Verdejo, G. Macia-Fernandez, and E. Vazquez, "Anomaly-based network intrusion detection: Techniques, systems and challenges," Computers & Security, vol. 28, no. 1, pp. 18-28, 2009.',
        '[10] M. Ring, S. Wunderlich, D. Scheuring, D. Landes, and A. Hotho, "A Survey of Network-based Intrusion Detection Data Sets," Computers & Security, vol. 86, pp. 147-167, 2019.',
        '[11] P. Velickovic, G. Cucurull, A. Casanova, A. Romero, P. Lio, and Y. Bengio, "Graph Attention Networks," in Proc. ICLR, 2018.',
        '[12] I. Sharafaldin, A. H. Lashkari, and A. A. Ghorbani, "Toward Generating a New Intrusion Detection Dataset and Intrusion Traffic Characterization," in Proc. ICISSP, 2018.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        _pf(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, left_cm=0.5, first_indent_cm=-0.5,
            space_before=0, space_after=4, line=1.5)
        _run(p, ref, size=13)

# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    print("Generating BaoCao_C2GNN_Final.docx ...")
    doc = setup_document()

    page_cover(doc)
    page_gvhd(doc)
    page_phan_cong(doc)
    page_cam_on(doc)
    page_tom_tat(doc)
    page_muc_luc(doc)
    page_danh_sach_hinh(doc)
    page_danh_sach_bang(doc)
    page_tu_viet_tat(doc)

    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    ket_luan(doc)
    tai_lieu_tham_khao(doc)

    doc.save(str(OUT))
    print(f"  OK  {OUT}")
    print()
    print("Mo file trong Word, nhan Ctrl+A -> F9 de cap nhat muc luc va danh sach hinh/bang.")


if __name__ == "__main__":
    main()
